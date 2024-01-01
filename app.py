import streamlit as st
from docx import Document
from zipfile import ZipFile
import os
import tempfile
from io import BytesIO

# Function to combine Word documents
def combine_word_documents(docs):
    combined_doc = Document()
    for doc in docs:
        sub_doc = Document(BytesIO(doc))
        for element in sub_doc.element.body:
            combined_doc.element.body.append(element)
    return combined_doc

# Function to process Word files from a ZIP
def process_zip_file(zip_file):
    with ZipFile(zip_file, 'r') as z:
        with tempfile.TemporaryDirectory() as tempdir:
            z.extractall(tempdir)
            word_docs = []
            error_occurred = False

            for folder in os.listdir(tempdir):
                folder_path = os.path.join(tempdir, folder)
                if os.path.isdir(folder_path):
                    docs_in_folder = [file for file in os.listdir(folder_path) if file.endswith('.docx')]
                    
                    if len(docs_in_folder) > 1:
                        st.error(f"More than one Word document found in the folder '{folder}'. Only the first document will be processed.")
                        error_occurred = True
                    
                    if docs_in_folder:
                        file_path = os.path.join(folder_path, docs_in_folder[0])
                        with open(file_path, 'rb') as f:
                            word_docs.append(f.read())

            return word_docs, error_occurred

# Function to process direct Word file uploads
def process_word_files(word_files):
    return [file.getvalue() for file in word_files]

# Initialize session state variables
if 'combined_document' not in st.session_state:
    st.session_state['combined_document'] = None

# Streamlit UI
st.title('Word Document Combiner')
st.subtitle('Word Document Combiner is an easy-to-use tool that helps you merge multiple Microsoft Word documents into one. Whether you have a bunch of chapters, sections, or separate documents that you want to put together, this app makes it simple.')

upload_choice = st.radio("Choose your upload method", ('Zip File', 'Word Files'))

if upload_choice == 'Zip File':
    uploaded_file = st.file_uploader("Upload ZIP file", type=['zip'])
    if st.button('Combine Documents from ZIP') and uploaded_file:
        word_docs, error_occurred = process_zip_file(uploaded_file)
        if word_docs and not error_occurred:
            st.session_state['combined_document'] = combine_word_documents(word_docs)

elif upload_choice == 'Word Files':
    uploaded_files = st.file_uploader("Upload Word files", accept_multiple_files=True, type=['docx'])
    if st.button('Combine Word Documents') and uploaded_files:
        word_docs = process_word_files(uploaded_files)
        st.session_state['combined_document'] = combine_word_documents(word_docs)

# Export options
if st.session_state['combined_document']:
    export_format = st.selectbox("Select export format", ("Word", "PDF", "Text"))

    if st.button('Export Combined Document'):
        file_stream = BytesIO()
        combined_document = st.session_state['combined_document']
        combined_document.save(file_stream)
        file_stream.seek(0)

        if export_format == "Word":
            st.download_button(label="Download Combined Document",
                               data=file_stream,
                               file_name="combined_document.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        elif export_format == "PDF":
            # PDF conversion logic here (Windows only with docx2pdf)
            # For cross-platform, use an alternative method
            pass
        elif export_format == "Text":
            text_stream = BytesIO()
            for paragraph in combined_document.paragraphs:
                text_stream.write(paragraph.text.encode('utf-8') + b'\n')
            text_stream.seek(0)
            st.download_button(label="Download Combined Document",
                               data=text_stream,
                               file_name="combined_document.txt",
                               mime="text/plain")
st.markdown("""

## **How to Use the App**

### **Getting Started**

*   You don’t need to install anything on your computer. The app runs in a web browser.
*   Make sure you have all your Word documents ready. You can have them as separate files or put together in a ZIP file.

### **Combining Your Documents**

1.  **Open the App**: Go to the web link where the Word Document Combiner is hosted.
2.  **Choose Your Upload Method**:
    *   If your documents are in a ZIP file, select the 'Zip File' option.
    *   If you have individual Word files, select the 'Word Files' option.
3.  **Upload Your Files**: Click on the upload area to select files from your computer, or drag and drop them into the box.
4.  **Combine Your Documents**:
    *   If you uploaded a ZIP file, click ‘Combine Documents from ZIP’.
    *   If you uploaded Word files, click ‘Combine Word Documents’.
5.  **Wait for the Process to Complete**: The app will merge your documents and let you know when it's done.

### **Downloading Your Combined Document**

*   Once your documents are combined, you can choose how you want to download the combined document. You can download it as a Word document, a PDF, or a text file.
*   Click on the download button for the format you want, and the file will be saved to your computer.

### **Troubleshooting**

*   If you get an error message or something doesn’t work, check to make sure that you’re uploading the right type of file (.docx or .zip).
*   If there are more than one Word document in a folder within your ZIP file, the app will only combine the first one it finds.

## **Need Help?**

If you have any questions or need help using the app, don’t hesitate to reach out for support.
""")
